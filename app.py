import streamlit as st
from openai import OpenAI
import os
from dotenv import load_dotenv
from io import BytesIO
import docx
import fitz  # PyMuPDF
from datetime import datetime
import json
import uuid

load_dotenv()

client = OpenAI(
    api_key=os.getenv("XAI_API_KEY"),
    base_url="https://api.x.ai/v1",
)

st.set_page_config(page_title="CV Tailor", page_icon="📄", layout="centered")

st.title("📄 CV Tailor for Mum 💕")
st.write("Upload → Choose style → Get tailored CV + Cover Letter")

# ====================== HISTORY ======================
HISTORY_FILE = "tailoring_history.json"

def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r") as f:
            return json.load(f)
    return []

def save_history(entry):
    history = load_history()
    history.append(entry)
    with open(HISTORY_FILE, "w") as f:
        json.dump(history, f, indent=2)

history = load_history()

# ====================== SIDEBAR ======================
with st.sidebar:
    st.header("📚 History")
    if history:
        for i, entry in enumerate(reversed(history[-10:])):  # Show last 10
            with st.expander(f"{entry['date']} - {entry['job_title'][:40]}..."):
                st.write(f"**Style:** {entry['style']}")
                if st.button("Download again", key=f"dl_{i}"):
                    bio = BytesIO()
                    doc = docx.Document(BytesIO(entry['docx_bytes']))
                    doc.save(bio)
                    bio.seek(0)
                    st.download_button("Save .docx", bio.getvalue(), 
                                     f"{entry['job_title']}.docx", key=f"dlbtn_{i}")
    else:
        st.info("No previous jobs yet.")

# ====================== MAIN APP ======================
cv_file = st.file_uploader("Your Current CV (PDF or DOCX)", type=["pdf", "docx"])
jd_file = st.file_uploader("Job Description (PDF, DOCX or TXT)", type=["pdf", "docx", "txt"])

col1, col2 = st.columns(2)
with col1:
    style = st.selectbox("CV Style", 
        ["Classic (Safe & Professional)", 
         "Modern (Clean & Contemporary)", 
         "Compact (Shorter & Punchy)", 
         "Achievement-Focused"])
with col2:
    emphasis = st.selectbox("Emphasize", ["Years of Experience", "Key Achievements", "Skills & Tools"])

if st.button("✨ Tailor with Grok", type="primary", use_container_width=True):
    if not cv_file or not jd_file:
        st.error("Please upload both files")
    else:
        with st.spinner("Reading files and asking Grok..."):
            # Extract text
            def extract_text(file):
                if file.type == "application/pdf":
                    text = ""
                    doc = fitz.open(stream=file.read(), filetype="pdf")
                    for page in doc:
                        text += page.get_text()
                    return text
                else:
                    doc = docx.Document(file)
                    return "\n".join([p.text for p in doc.paragraphs])

            cv_text = extract_text(cv_file)
            if jd_file.type == "text/plain":
                jd_text = jd_file.read().decode("utf-8")
            else:
                jd_text = extract_text(jd_file)

            # Style instructions
            style_instructions = {
                "Classic (Safe & Professional)": "Use traditional CV format with clear sections, bullet points, and professional language.",
                "Modern (Clean & Contemporary)": "Use modern layout with bold section headers, two-column skills if possible, and contemporary language.",
                "Compact (Shorter & Punchy)": "Keep it to 1-2 pages max. Use strong action verbs and very concise bullets.",
                "Achievement-Focused": "Emphasize quantifiable achievements. Start every bullet with strong action verbs and numbers."
            }

            prompt = f"""You are an expert careers advisor.

**Task**: Create a perfectly tailored CV and cover letter for the job.

**Style**: {style_instructions[style]}
**Emphasis**: {emphasis}

**Original CV**:
{cv_text}

**Job Description**:
{jd_text}

Output in clean Markdown:
1. **TAILORED CV** (full version with all sections)
2. **COVER LETTER** (3-4 strong paragraphs)

Use keywords from the job description naturally. Make it ATS-friendly."""

            response = client.chat.completions.create(
                model="grok-4",   # or grok-3 if preferred
                messages=[
                    {"role": "system", "content": "You are a senior HR professional and careers coach."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.65,
                max_tokens=5000
            )

            result = response.choices[0].message.content

        # Create Word Document with formatting
        doc = docx.Document()
        doc.add_heading("Tailored CV & Cover Letter", 0)
        doc.add_paragraph(f"Job: {jd_text[:100]}... | Style: {style} | Date: {datetime.now().strftime('%d %b %Y')}")
        doc.add_paragraph("")

        # Simple formatting
        for line in result.split("\n"):
            if line.startswith("# ") or line.startswith("## "):
                doc.add_heading(line.replace("#", "").strip(), level=1 if line.startswith("# ") else 2)
            elif line.strip().startswith("- ") or line.strip().startswith("•"):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(line)
        
        # Save to memory for history
        bio = BytesIO()
        doc.save(bio)
        docx_bytes = bio.getvalue()

        # Save to history
        job_title = jd_text.split("\n")[0][:60] if jd_text else "Unknown Role"
        history_entry = {
            "id": str(uuid.uuid4()),
            "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "job_title": job_title,
            "style": style,
            "result": result,
            "docx_bytes": list(docx_bytes)  # JSON serializable
        }
        save_history(history_entry)

        st.success("✅ Tailored successfully!")

        st.subheader("Preview")
        st.markdown(result)

        # Download button
        st.download_button(
            label="📥 Download as Word Document (.docx)",
            data=docx_bytes,
            file_name=f"CV_{job_title[:30]}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        st.info("💡 Open the .docx in Microsoft Word or upload to Google Docs (File → Import) to edit further.")